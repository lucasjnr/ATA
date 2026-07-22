"""
Knowledge Base document extractor.
Extracts text, styles, linguistic patterns, and structural features
from PDF, DOCX, XLSX, and TXT files.
"""
import io
import re
import hashlib
import logging
from typing import Optional

logger = logging.getLogger(__name__)

# ── Formal linguistic markers used in RPPS documents ──
FORMAL_MARKERS = [
    "nada mais havendo a tratar",
    "comitê de investimentos",
    "deliberou",
    "deliberação",
    "decisão",
    "encaminhamento",
    "lavrou",
    "secretário",
    "presidente",
    "reunião ordinária",
    "reunião extraordinária",
    "quórum",
    "pauta",
    "regime próprio",
    "rpps",
    "ata de reunião",
    "conselho deliberativo",
    "conselho fiscal",
    "diretoria executiva",
]

THIRD_PERSON_INDICATORS = ["sr.", "sra.", "o conselho", "a diretoria", "os membros", "foi deliberado", "foram aprovados"]
MODAL_VERBS = ["deve", "deverá", "poderá", "recomenda-se", "estabelece-se", "determina-se"]


def _extract_text_pdf(data: bytes) -> str:
    """Extract text from PDF bytes."""
    try:
        import pypdf
        reader = pypdf.PdfReader(io.BytesIO(data))
        pages = []
        for page in reader.pages:
            text = page.extract_text() or ""
            pages.append(text)
        return "\n".join(pages)
    except ImportError:
        pass
    try:
        # Fallback: minimal raw text extraction
        text = data.decode("latin-1", errors="ignore")
        lines = [l.strip() for l in text.split("\n") if l.strip() and len(l.strip()) > 3]
        return " ".join(lines[:200])
    except Exception as e:
        logger.warning(f"PDF extraction failed: {e}")
        return ""


def _extract_text_docx(data: bytes) -> str:
    """Extract text from DOCX bytes."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(data))
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    except Exception as e:
        logger.warning(f"DOCX extraction failed: {e}")
        return ""


def _extract_styles_docx(data: bytes) -> dict:
    """Extract font and style info from DOCX."""
    styles = {"fonts": [], "has_bold_headings": False, "has_numbered_sections": False}
    try:
        from docx import Document
        doc = Document(io.BytesIO(data))
        fonts_seen = set()
        for para in doc.paragraphs:
            for run in para.runs:
                if run.font.name and run.font.name not in fonts_seen:
                    fonts_seen.add(run.font.name)
                    styles["fonts"].append(run.font.name)
            if para.style.name.startswith("Heading") and any(r.bold for r in para.runs):
                styles["has_bold_headings"] = True
            if re.match(r"^\d+[\.\)]\s", para.text):
                styles["has_numbered_sections"] = True
    except Exception:
        pass
    return styles


def _extract_text_xlsx(data: bytes) -> str:
    """Extract text from XLSX bytes."""
    try:
        import openpyxl
        wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
        rows = []
        for sheet in wb.worksheets:
            for row in sheet.iter_rows(values_only=True):
                cells = [str(c) for c in row if c is not None and str(c).strip()]
                if cells:
                    rows.append(" | ".join(cells))
        return "\n".join(rows[:100])
    except Exception as e:
        logger.warning(f"XLSX extraction failed: {e}")
        return ""


def _extract_linguistic_patterns(text: str) -> dict:
    """Analyse text for formality and RPPS-specific linguistic patterns."""
    text_lower = text.lower()
    found_markers = [m for m in FORMAL_MARKERS if m in text_lower]
    third_person = sum(1 for t in THIRD_PERSON_INDICATORS if t in text_lower)
    modals = sum(1 for m in MODAL_VERBS if m in text_lower)
    word_count = len(text.split())

    formality_score = min(100, (len(found_markers) * 5) + (third_person * 3) + (modals * 2))
    avg_sentence_len = 0
    sentences = re.split(r"[.!?]+", text)
    sentences = [s.strip() for s in sentences if s.strip()]
    if sentences:
        avg_sentence_len = sum(len(s.split()) for s in sentences) / len(sentences)

    return {
        "formal_markers_found": found_markers,
        "formal_marker_count": len(found_markers),
        "third_person_count": third_person,
        "modal_verb_count": modals,
        "estimated_formality": formality_score,
        "word_count": word_count,
        "sentence_count": len(sentences),
        "avg_sentence_length": round(avg_sentence_len, 1),
    }


def _extract_structure(text: str) -> dict:
    """Identify document structure patterns."""
    lines = text.split("\n")
    headings = []
    numbered = False

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        if re.match(r"^\d+[\.\)]\s+\w", stripped) and len(stripped) < 120:
            headings.append(stripped[:80])
            numbered = True
        elif stripped.isupper() and 5 < len(stripped) < 100:
            headings.append(stripped[:80])

    common_sections = []
    keywords = ["abertura", "presentes", "pauta", "ordem do dia", "deliberações", "encerramento", "encaminhamentos", "próxima reunião"]
    for kw in keywords:
        if kw in text.lower():
            common_sections.append(kw)

    return {
        "detected_headings": headings[:20],
        "uses_numbered_sections": numbered,
        "common_sections": common_sections,
        "section_count": len(headings),
    }


def compute_sha1(data: bytes) -> str:
    return hashlib.sha1(data).hexdigest()


def extract_document_features(
    filename: str,
    filetype: str,
    data: bytes,
) -> dict:
    """
    Main entry point. Returns full feature dict for storage in knowledge_docs.metadata.
    """
    ft = filetype.lower()
    text = ""
    styles = {}

    if ft == "pdf":
        text = _extract_text_pdf(data)
    elif ft in ("docx", "doc"):
        text = _extract_text_docx(data)
        styles = _extract_styles_docx(data)
    elif ft in ("xlsx", "xls", "csv"):
        text = _extract_text_xlsx(data)
    elif ft == "txt":
        text = data.decode("utf-8", errors="ignore")
    else:
        try:
            text = data.decode("utf-8", errors="ignore")
        except Exception:
            text = ""

    linguistic = _extract_linguistic_patterns(text)
    structure = _extract_structure(text)

    # Brasão heuristic: small file (<200KB) in first-page context or named with keywords
    is_brasao_candidate = (
        len(data) < 200 * 1024
        and any(kw in filename.lower() for kw in ["brasao", "escudo", "logo", "braso", "simbolo"])
    )

    metadata = {
        "sha1": compute_sha1(data),
        "file_size_bytes": len(data),
        "filetype": ft,
        "is_brasao_candidate": is_brasao_candidate,
        "styles": styles,
        "linguistic_patterns": linguistic,
        "structure": structure,
        "text_preview": text[:500],
    }

    return text, metadata
