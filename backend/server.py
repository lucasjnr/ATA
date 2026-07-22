from fastapi import FastAPI, APIRouter, HTTPException, Depends, UploadFile, File, Form
from fastapi.responses import StreamingResponse, FileResponse
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
import os
import io
import logging
import uuid
import json
import bcrypt
import jwt as pyjwt
from pathlib import Path
from pydantic import BaseModel, Field, EmailStr
from typing import List, Optional, Literal
from datetime import datetime, timezone, timedelta

from emergentintegrations.llm.openai import OpenAISpeechToText
from emergentintegrations.llm.chat import LlmChat, UserMessage

from services.knowledge_extractor import extract_document_features
from services.scoring_engine import score_meeting, detect_gaps
from services.economic_data import get_economic_indicators, get_selic_history, get_ipca_history, generate_scenario_analysis
import services.intelligent_agent as agent_svc

from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.enums import TA_JUSTIFY
from docx import Document as DocxDocument
from docx.shared import Pt

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

mongo_url = os.environ['MONGO_URL']
client = AsyncIOMotorClient(mongo_url)
db = client[os.environ['DB_NAME']]

JWT_SECRET = os.environ['JWT_SECRET']
JWT_ALGORITHM = os.environ['JWT_ALGORITHM']
JWT_EXPIRY_HOURS = int(os.environ['JWT_EXPIRY_HOURS'])
EMERGENT_LLM_KEY = os.environ['EMERGENT_LLM_KEY']

agent_svc.EMERGENT_LLM_KEY = EMERGENT_LLM_KEY

app = FastAPI(title="Sistema de Atas de Reunião IA")
api_router = APIRouter(prefix="/api")
security = HTTPBearer()

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)


# =========================
# Models
# =========================
def now_iso():
    return datetime.now(timezone.utc).isoformat()


class UserRegister(BaseModel):
    name: str
    email: EmailStr
    password: str
    institution: Optional[str] = ""


class UserLogin(BaseModel):
    email: EmailStr
    password: str


class User(BaseModel):
    id: str
    name: str
    email: str
    institution: str = ""
    role: str = "admin"
    created_at: str


class Participant(BaseModel):
    name: str
    role: str = ""
    email: str = ""
    institution: str = ""
    present: bool = True


class AgendaTopic(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    title: str
    description: str = ""
    order: int = 0
    status: Literal["pending", "in_progress", "completed"] = "pending"
    transcription: str = ""
    summary: str = ""
    deliberation: str = ""
    encaminhamento: str = ""
    responsible: str = ""
    deadline: str = ""


class MeetingCreate(BaseModel):
    title: str
    meeting_type: str = "Reunião Administrativa"
    date: str
    start_time: str
    end_time: str = ""
    location: str = ""
    responsible: str = ""
    secretary: str = ""
    president: str = ""
    participants: List[Participant] = []
    agenda: List[AgendaTopic] = []


class MeetingUpdate(BaseModel):
    title: Optional[str] = None
    meeting_type: Optional[str] = None
    date: Optional[str] = None
    start_time: Optional[str] = None
    end_time: Optional[str] = None
    location: Optional[str] = None
    responsible: Optional[str] = None
    secretary: Optional[str] = None
    president: Optional[str] = None
    participants: Optional[List[Participant]] = None
    agenda: Optional[List[AgendaTopic]] = None
    status: Optional[str] = None
    ata_content: Optional[str] = None
    executive_summary: Optional[str] = None
    next_agenda: Optional[str] = None


class Meeting(BaseModel):
    id: str
    number: int
    user_id: str
    title: str
    meeting_type: str
    date: str
    start_time: str
    end_time: str
    location: str
    responsible: str
    secretary: str
    president: str
    participants: List[Participant]
    agenda: List[AgendaTopic]
    status: str = "scheduled"
    ata_content: str = ""
    executive_summary: str = ""
    next_agenda: str = ""
    created_at: str
    updated_at: str


class Deliberation(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    meeting_id: str
    meeting_number: int
    topic_title: str
    decision: str
    responsible: str = ""
    deadline: str = ""
    status: Literal["aberto", "em_andamento", "concluido", "cancelado"] = "aberto"
    created_at: str = Field(default_factory=now_iso)


# =========================
# Auth helpers
# =========================
def hash_pw(pw: str) -> str:
    return bcrypt.hashpw(pw.encode(), bcrypt.gensalt()).decode()


def verify_pw(pw: str, hashed: str) -> bool:
    return bcrypt.checkpw(pw.encode(), hashed.encode())


def create_token(user_id: str) -> str:
    payload = {
        "sub": user_id,
        "exp": datetime.now(timezone.utc) + timedelta(hours=JWT_EXPIRY_HOURS),
        "iat": datetime.now(timezone.utc),
    }
    return pyjwt.encode(payload, JWT_SECRET, algorithm=JWT_ALGORITHM)


async def get_current_user(credentials: HTTPAuthorizationCredentials = Depends(security)) -> dict:
    try:
        payload = pyjwt.decode(credentials.credentials, JWT_SECRET, algorithms=[JWT_ALGORITHM])
        user_id = payload.get("sub")
    except pyjwt.PyJWTError:
        raise HTTPException(status_code=401, detail="Token inválido")
    user = await db.users.find_one({"id": user_id}, {"_id": 0, "password": 0})
    if not user:
        raise HTTPException(status_code=401, detail="Usuário não encontrado")
    return user


# =========================
# Auth Routes
# =========================
@api_router.post("/auth/register")
async def register(payload: UserRegister):
    existing = await db.users.find_one({"email": payload.email.lower()})
    if existing:
        raise HTTPException(status_code=400, detail="E-mail já cadastrado")
    user_doc = {
        "id": str(uuid.uuid4()),
        "name": payload.name,
        "email": payload.email.lower(),
        "password": hash_pw(payload.password),
        "institution": payload.institution or "",
        "role": "admin",
        "created_at": now_iso(),
    }
    await db.users.insert_one(user_doc)
    token = create_token(user_doc["id"])
    user_doc.pop("password")
    user_doc.pop("_id", None)
    return {"token": token, "user": user_doc}


@api_router.post("/auth/login")
async def login(payload: UserLogin):
    user = await db.users.find_one({"email": payload.email.lower()})
    if not user or not verify_pw(payload.password, user["password"]):
        raise HTTPException(status_code=401, detail="Credenciais inválidas")
    token = create_token(user["id"])
    user.pop("password")
    user.pop("_id", None)
    return {"token": token, "user": user}


@api_router.get("/auth/me")
async def me(user=Depends(get_current_user)):
    return user


# =========================
# Meetings
# =========================
async def next_meeting_number(user_id: str) -> int:
    last = await db.meetings.find({"user_id": user_id}).sort("number", -1).limit(1).to_list(1)
    return (last[0]["number"] + 1) if last else 1


@api_router.post("/meetings")
async def create_meeting(payload: MeetingCreate, user=Depends(get_current_user)):
    number = await next_meeting_number(user["id"])
    meeting = {
        "id": str(uuid.uuid4()),
        "number": number,
        "user_id": user["id"],
        "title": payload.title,
        "meeting_type": payload.meeting_type,
        "date": payload.date,
        "start_time": payload.start_time,
        "end_time": payload.end_time,
        "location": payload.location,
        "responsible": payload.responsible,
        "secretary": payload.secretary,
        "president": payload.president,
        "participants": [p.model_dump() for p in payload.participants],
        "agenda": [a.model_dump() for a in payload.agenda],
        "status": "scheduled",
        "ata_content": "",
        "executive_summary": "",
        "next_agenda": "",
        "created_at": now_iso(),
        "updated_at": now_iso(),
    }
    await db.meetings.insert_one(meeting)
    meeting.pop("_id", None)
    return meeting


@api_router.get("/meetings")
async def list_meetings(user=Depends(get_current_user), status: Optional[str] = None, search: Optional[str] = None):
    query = {"user_id": user["id"]}
    if status:
        query["status"] = status
    if search:
        query["$or"] = [
            {"title": {"$regex": search, "$options": "i"}},
            {"meeting_type": {"$regex": search, "$options": "i"}},
            {"location": {"$regex": search, "$options": "i"}},
        ]
    docs = await db.meetings.find(query, {"_id": 0}).sort("date", -1).to_list(500)
    return docs


@api_router.get("/meetings/{meeting_id}")
async def get_meeting(meeting_id: str, user=Depends(get_current_user)):
    doc = await db.meetings.find_one({"id": meeting_id, "user_id": user["id"]}, {"_id": 0})
    if not doc:
        raise HTTPException(status_code=404, detail="Reunião não encontrada")
    return doc


@api_router.patch("/meetings/{meeting_id}")
async def update_meeting(meeting_id: str, payload: MeetingUpdate, user=Depends(get_current_user)):
    update = {k: v for k, v in payload.model_dump(exclude_none=True).items()}
    if "participants" in update:
        update["participants"] = [p if isinstance(p, dict) else p.model_dump() for p in update["participants"]]
    if "agenda" in update:
        update["agenda"] = [a if isinstance(a, dict) else a.model_dump() for a in update["agenda"]]
    update["updated_at"] = now_iso()
    res = await db.meetings.find_one_and_update(
        {"id": meeting_id, "user_id": user["id"]},
        {"$set": update},
        return_document=True,
        projection={"_id": 0},
    )
    if not res:
        raise HTTPException(status_code=404, detail="Reunião não encontrada")
    return res


@api_router.delete("/meetings/{meeting_id}")
async def delete_meeting(meeting_id: str, user=Depends(get_current_user)):
    await db.meetings.delete_one({"id": meeting_id, "user_id": user["id"]})
    await db.deliberations.delete_many({"meeting_id": meeting_id})
    return {"ok": True}


@api_router.post("/meetings/{meeting_id}/start")
async def start_meeting(meeting_id: str, user=Depends(get_current_user)):
    res = await db.meetings.find_one_and_update(
        {"id": meeting_id, "user_id": user["id"]},
        {"$set": {"status": "in_progress", "updated_at": now_iso()}},
        return_document=True,
        projection={"_id": 0},
    )
    if not res:
        raise HTTPException(status_code=404, detail="Reunião não encontrada")
    return res


# =========================
# Transcription / AI
# =========================
@api_router.post("/meetings/{meeting_id}/topics/{topic_id}/transcribe")
async def transcribe_topic(
    meeting_id: str,
    topic_id: str,
    file: UploadFile = File(...),
    user=Depends(get_current_user),
):
    meeting = await db.meetings.find_one({"id": meeting_id, "user_id": user["id"]})
    if not meeting:
        raise HTTPException(status_code=404, detail="Reunião não encontrada")

    audio_bytes = await file.read()
    if len(audio_bytes) > 25 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="Arquivo de áudio excede 25MB")

    stt = OpenAISpeechToText(api_key=EMERGENT_LLM_KEY)
    audio_io = io.BytesIO(audio_bytes)
    audio_io.name = file.filename or "audio.webm"
    try:
        response = await stt.transcribe(
            file=audio_io,
            model="whisper-1",
            response_format="json",
            language="pt",
        )
        text = response.text
    except Exception as e:
        logger.exception("Erro na transcrição")
        raise HTTPException(status_code=500, detail=f"Erro na transcrição: {e}")

    agenda = meeting.get("agenda", [])
    for t in agenda:
        if t["id"] == topic_id:
            t["transcription"] = (t.get("transcription", "") + " " + text).strip()
            break
    await db.meetings.update_one(
        {"id": meeting_id},
        {"$set": {"agenda": agenda, "updated_at": now_iso()}},
    )
    return {"text": text}


@api_router.post("/meetings/{meeting_id}/topics/{topic_id}/finalize")
async def finalize_topic(meeting_id: str, topic_id: str, user=Depends(get_current_user)):
    meeting = await db.meetings.find_one({"id": meeting_id, "user_id": user["id"]}, {"_id": 0})
    if not meeting:
        raise HTTPException(status_code=404, detail="Reunião não encontrada")

    topic = next((t for t in meeting.get("agenda", []) if t["id"] == topic_id), None)
    if not topic:
        raise HTTPException(status_code=404, detail="Tópico não encontrado")

    transcription = topic.get("transcription", "").strip()
    if not transcription:
        transcription = topic.get("summary") or topic.get("title", "")

    prompt = f"""Você é um secretário especialista em atas de reuniões institucionais brasileiras.
Analise a transcrição abaixo do tópico "{topic['title']}" da reunião e gere um JSON com a estrutura:
{{
  "discussao": "Resumo objetivo da discussão em 2-4 frases, em português formal.",
  "deliberacao": "Decisão tomada ou consenso, em uma frase clara. Se não houve, escreva 'Sem deliberação registrada.'",
  "encaminhamento": "Ação prática a ser realizada. Se não houver, 'Sem encaminhamento.'",
  "responsavel": "Nome do responsável citado ou 'Não definido'",
  "prazo": "Prazo mencionado (formato livre, ex. '30/03/2026' ou '30 dias') ou 'Não definido'"
}}

Transcrição:
\"\"\"{transcription}\"\"\"

Responda APENAS com o JSON, sem markdown ou explicações."""

    chat = LlmChat(
        api_key=EMERGENT_LLM_KEY,
        session_id=f"finalize-{meeting_id}-{topic_id}",
        system_message="Você é um assistente especialista em redação de atas institucionais em português do Brasil. Sempre retorna JSON válido.",
    ).with_model("openai", "gpt-5.2")

    try:
        result = await chat.send_message(UserMessage(text=prompt))
        content = result.strip()
        if content.startswith("```"):
            content = content.split("```")[1]
            if content.startswith("json"):
                content = content[4:]
            content = content.strip()
        data = json.loads(content)
    except Exception:
        logger.exception("Erro IA finalize")
        data = {
            "discussao": transcription[:500],
            "deliberacao": "Sem deliberação registrada.",
            "encaminhamento": "Sem encaminhamento.",
            "responsavel": "Não definido",
            "prazo": "Não definido",
        }

    agenda = meeting["agenda"]
    for t in agenda:
        if t["id"] == topic_id:
            t["summary"] = data.get("discussao", "")
            t["deliberation"] = data.get("deliberacao", "")
            t["encaminhamento"] = data.get("encaminhamento", "")
            t["responsible"] = data.get("responsavel", "")
            t["deadline"] = data.get("prazo", "")
            t["status"] = "completed"
            break

    await db.meetings.update_one(
        {"id": meeting_id}, {"$set": {"agenda": agenda, "updated_at": now_iso()}}
    )

    if data.get("deliberacao") and "sem deliberação" not in data["deliberacao"].lower():
        delib = {
            "id": str(uuid.uuid4()),
            "meeting_id": meeting_id,
            "meeting_number": meeting["number"],
            "topic_title": topic["title"],
            "decision": data.get("deliberacao", ""),
            "responsible": data.get("responsavel", ""),
            "deadline": data.get("prazo", ""),
            "status": "aberto",
            "user_id": user["id"],
            "created_at": now_iso(),
        }
        await db.deliberations.insert_one(delib)

    return data


@api_router.post("/meetings/{meeting_id}/generate-ata")
async def generate_ata(meeting_id: str, user=Depends(get_current_user)):
    meeting = await db.meetings.find_one({"id": meeting_id, "user_id": user["id"]}, {"_id": 0})
    if not meeting:
        raise HTTPException(status_code=404, detail="Reunião não encontrada")

    agenda_summary = "\n\n".join(
        f"### {i+1}. {t['title']}\n"
        f"Discussão: {t.get('summary', '')}\n"
        f"Deliberação: {t.get('deliberation', '')}\n"
        f"Encaminhamento: {t.get('encaminhamento', '')}\n"
        f"Responsável: {t.get('responsible', '')}  |  Prazo: {t.get('deadline', '')}"
        for i, t in enumerate(meeting.get("agenda", []))
    )
    participants_present = ", ".join(p["name"] for p in meeting.get("participants", []) if p.get("present"))
    participants_absent = ", ".join(p["name"] for p in meeting.get("participants", []) if not p.get("present"))

    prompt = f"""Você é um secretário oficial. Gere uma ATA DE REUNIÃO formal, em português do Brasil, com tom institucional.

Dados:
- Número da reunião: {meeting['number']}
- Tipo: {meeting['meeting_type']}
- Título: {meeting['title']}
- Data: {meeting['date']}
- Horário: {meeting['start_time']} às {meeting.get('end_time', '—')}
- Local: {meeting['location']}
- Presidente/Coordenador: {meeting.get('president', '—')}
- Secretário: {meeting.get('secretary', '—')}
- Presentes: {participants_present or '—'}
- Ausentes: {participants_absent or '—'}

Pauta e deliberações:
{agenda_summary}

Retorne JSON válido:
{{
  "ata_completa": "Texto completo da ata em HTML simples (use <h2>, <h3>, <p>, <ul>, <li>, <strong>). Inclua cabeçalho, presentes, ordem do dia, discussões item a item, deliberações, encerramento.",
  "resumo_executivo": "Resumo executivo em 1 parágrafo curto (até 150 palavras).",
  "proxima_pauta_sugerida": "Lista em HTML <ul><li>...</li></ul> de 3-5 tópicos sugeridos para a próxima reunião com base nas pendências."
}}

Responda APENAS com o JSON puro."""

    chat = LlmChat(
        api_key=EMERGENT_LLM_KEY,
        session_id=f"ata-{meeting_id}",
        system_message="Você redige atas oficiais em português brasileiro com precisão jurídica e clareza.",
    ).with_model("openai", "gpt-5.2")

    try:
        result = await chat.send_message(UserMessage(text=prompt))
        content = result.strip()
        if content.startswith("```"):
            content = content.split("```")[1]
            if content.startswith("json"):
                content = content[4:]
            content = content.strip()
        data = json.loads(content)
    except Exception as e:
        logger.exception("Erro IA gerar ata")
        raise HTTPException(status_code=500, detail=f"Erro ao gerar ata: {e}")

    await db.meetings.update_one(
        {"id": meeting_id},
        {"$set": {
            "ata_content": data.get("ata_completa", ""),
            "executive_summary": data.get("resumo_executivo", ""),
            "next_agenda": data.get("proxima_pauta_sugerida", ""),
            "status": "completed",
            "updated_at": now_iso(),
        }},
    )
    return data


# =========================
# Export
# =========================
def html_to_text(html: str) -> str:
    import re
    txt = re.sub(r"<br\s*/?>", "\n", html)
    txt = re.sub(r"</p>", "\n\n", txt)
    txt = re.sub(r"</h\d>", "\n\n", txt)
    txt = re.sub(r"<li>", "• ", txt)
    txt = re.sub(r"</li>", "\n", txt)
    txt = re.sub(r"<[^>]+>", "", txt)
    return txt.strip()


@api_router.get("/meetings/{meeting_id}/export/pdf")
async def export_pdf(meeting_id: str, user=Depends(get_current_user)):
    meeting = await db.meetings.find_one({"id": meeting_id, "user_id": user["id"]}, {"_id": 0})
    if not meeting:
        raise HTTPException(status_code=404, detail="Reunião não encontrada")

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, leftMargin=2 * cm, rightMargin=2 * cm, topMargin=2 * cm, bottomMargin=2 * cm)
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="Body", fontSize=11, leading=16, alignment=TA_JUSTIFY))
    story = []
    story.append(Paragraph(f"<b>ATA Nº {meeting['number']} — {meeting['title']}</b>", styles["Title"]))
    story.append(Spacer(1, 0.5 * cm))
    story.append(Paragraph(f"<b>Tipo:</b> {meeting['meeting_type']}", styles["Body"]))
    story.append(Paragraph(f"<b>Data:</b> {meeting['date']} — <b>Horário:</b> {meeting['start_time']} às {meeting.get('end_time','—')}", styles["Body"]))
    story.append(Paragraph(f"<b>Local:</b> {meeting['location']}", styles["Body"]))
    story.append(Spacer(1, 0.4 * cm))

    ata = meeting.get("ata_content") or "Ata ainda não gerada."
    for para in html_to_text(ata).split("\n\n"):
        if para.strip():
            story.append(Paragraph(para.replace("\n", "<br/>"), styles["Body"]))
            story.append(Spacer(1, 0.3 * cm))

    doc.build(story)
    buf.seek(0)
    return StreamingResponse(
        buf,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="ata_{meeting["number"]}.pdf"'},
    )


@api_router.get("/meetings/{meeting_id}/export/docx")
async def export_docx(meeting_id: str, user=Depends(get_current_user)):
    meeting = await db.meetings.find_one({"id": meeting_id, "user_id": user["id"]}, {"_id": 0})
    if not meeting:
        raise HTTPException(status_code=404, detail="Reunião não encontrada")

    docx = DocxDocument()
    style = docx.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    docx.add_heading(f"ATA Nº {meeting['number']} — {meeting['title']}", level=1)
    docx.add_paragraph(f"Tipo: {meeting['meeting_type']}")
    docx.add_paragraph(f"Data: {meeting['date']} | Horário: {meeting['start_time']} - {meeting.get('end_time','—')}")
    docx.add_paragraph(f"Local: {meeting['location']}")
    docx.add_paragraph("")

    ata_text = html_to_text(meeting.get("ata_content") or "Ata ainda não gerada.")
    for para in ata_text.split("\n\n"):
        if para.strip():
            docx.add_paragraph(para)

    buf = io.BytesIO()
    docx.save(buf)
    buf.seek(0)
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="ata_{meeting["number"]}.docx"'},
    )


# =========================
# Deliberations
# =========================
@api_router.get("/deliberations")
async def list_deliberations(user=Depends(get_current_user), status: Optional[str] = None):
    query = {"user_id": user["id"]}
    if status:
        query["status"] = status
    docs = await db.deliberations.find(query, {"_id": 0}).sort("created_at", -1).to_list(500)
    return docs


@api_router.patch("/deliberations/{deliberation_id}")
async def update_deliberation(deliberation_id: str, payload: dict, user=Depends(get_current_user)):
    res = await db.deliberations.find_one_and_update(
        {"id": deliberation_id, "user_id": user["id"]},
        {"$set": payload},
        return_document=True,
        projection={"_id": 0},
    )
    if not res:
        raise HTTPException(status_code=404, detail="Deliberação não encontrada")
    return res


# =========================
# Dashboard stats
# =========================
@api_router.get("/dashboard/stats")
async def dashboard_stats(user=Depends(get_current_user)):
    total = await db.meetings.count_documents({"user_id": user["id"]})
    completed = await db.meetings.count_documents({"user_id": user["id"], "status": "completed"})
    in_progress = await db.meetings.count_documents({"user_id": user["id"], "status": "in_progress"})
    scheduled = await db.meetings.count_documents({"user_id": user["id"], "status": "scheduled"})
    pending_delibs = await db.deliberations.count_documents({"user_id": user["id"], "status": {"$in": ["aberto", "em_andamento"]}})
    return {
        "total_meetings": total,
        "completed_meetings": completed,
        "in_progress_meetings": in_progress,
        "upcoming_meetings": scheduled,
        "pending_deliberations": pending_delibs,
    }


# =========================
# Templates
# =========================
DEFAULT_TEMPLATES = [
    {"id": "diretoria", "name": "Ata de Diretoria", "agenda": [
        {"title": "Aprovação da ata anterior"},
        {"title": "Apresentação de resultados"},
        {"title": "Planejamento de ações"},
        {"title": "Assuntos gerais"},
    ]},
    {"id": "conselho", "name": "Ata de Conselho", "agenda": [
        {"title": "Verificação de quórum"},
        {"title": "Leitura da pauta"},
        {"title": "Deliberações"},
        {"title": "Encerramento"},
    ]},
    {"id": "comissao", "name": "Ata de Comissão", "agenda": [
        {"title": "Abertura"},
        {"title": "Discussão dos temas"},
        {"title": "Encaminhamentos"},
    ]},
    {"id": "tecnica", "name": "Ata de Reunião Técnica", "agenda": [
        {"title": "Apresentação do tema técnico"},
        {"title": "Discussão"},
        {"title": "Próximos passos"},
    ]},
]


@api_router.get("/templates")
async def list_templates(user=Depends(get_current_user)):
    return DEFAULT_TEMPLATES


# =========================
# Knowledge Base
# =========================
@api_router.post("/kb/upload")
async def kb_upload(
    file: UploadFile = File(...),
    user=Depends(get_current_user),
):
    data = await file.read()
    if len(data) > 20 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="Arquivo excede 20MB")

    filename = file.filename or "document"
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else "txt"

    text, metadata = extract_document_features(filename, ext, data)

    doc = {
        "id": str(uuid.uuid4()),
        "user_id": user["id"],
        "filename": filename,
        "filetype": ext,
        "text": text,
        "metadata": metadata,
        "created_at": now_iso(),
    }
    existing = await db.knowledge_docs.find_one(
        {"user_id": user["id"], "metadata.sha1": metadata["sha1"]}
    )
    if existing:
        existing.pop("_id", None)
        return existing

    await db.knowledge_docs.insert_one(doc)
    doc.pop("_id", None)
    doc.pop("text", None)
    return doc


@api_router.get("/kb")
async def kb_list(user=Depends(get_current_user)):
    docs = await db.knowledge_docs.find(
        {"user_id": user["id"]}, {"_id": 0, "text": 0}
    ).sort("created_at", -1).to_list(200)
    return docs


@api_router.get("/kb/search")
async def kb_search(q: str, user=Depends(get_current_user)):
    if not q or len(q) < 2:
        return []
    docs = await db.knowledge_docs.find(
        {
            "user_id": user["id"],
            "$or": [
                {"filename": {"$regex": q, "$options": "i"}},
                {"text": {"$regex": q, "$options": "i"}},
            ],
        },
        {"_id": 0, "text": 0},
    ).limit(20).to_list(20)
    return docs


@api_router.get("/kb/{doc_id}")
async def kb_get(doc_id: str, user=Depends(get_current_user)):
    doc = await db.knowledge_docs.find_one(
        {"id": doc_id, "user_id": user["id"]}, {"_id": 0}
    )
    if not doc:
        raise HTTPException(status_code=404, detail="Documento não encontrado")
    doc.pop("text", None)
    return doc


@api_router.delete("/kb/{doc_id}")
async def kb_delete(doc_id: str, user=Depends(get_current_user)):
    await db.knowledge_docs.delete_one({"id": doc_id, "user_id": user["id"]})
    return {"ok": True}


# =========================
# Intelligent Agent
# =========================
@api_router.post("/meetings/{meeting_id}/agent/run")
async def agent_run(meeting_id: str, user=Depends(get_current_user)):
    meeting = await db.meetings.find_one({"id": meeting_id, "user_id": user["id"]}, {"_id": 0})
    if not meeting:
        raise HTTPException(status_code=404, detail="Reunião não encontrada")

    kb_docs = await db.knowledge_docs.find({"user_id": user["id"]}).to_list(50)
    settings_doc = await db.admin_settings.find_one({"user_id": user["id"]}) or {}
    settings_doc.pop("_id", None)

    try:
        result = await agent_svc.run_agent(
            meeting=meeting,
            kb_docs=kb_docs,
            settings=settings_doc,
            api_key=EMERGENT_LLM_KEY,
        )
    except Exception as e:
        logger.exception("Agent run error")
        raise HTTPException(status_code=500, detail=f"Erro no agente: {e}")

    await db.agent_changes.delete_many({"meeting_id": meeting_id})
    if result["changes"]:
        for ch in result["changes"]:
            ch["user_id"] = user["id"]
        await db.agent_changes.insert_many([dict(c) for c in result["changes"]])

    result.pop("changes", None)
    result["changes_count"] = result.get("changes_count", 0)
    return result


@api_router.get("/meetings/{meeting_id}/agent/changes")
async def agent_changes(meeting_id: str, user=Depends(get_current_user)):
    docs = await db.agent_changes.find(
        {"meeting_id": meeting_id, "user_id": user["id"]}, {"_id": 0}
    ).to_list(100)
    return docs


@api_router.get("/meetings/{meeting_id}/agent/pending")
async def agent_pending(meeting_id: str, user=Depends(get_current_user)):
    docs = await db.agent_changes.find(
        {"meeting_id": meeting_id, "user_id": user["id"], "status": "pending"}, {"_id": 0}
    ).to_list(100)
    return docs


class AgentFeedback(BaseModel):
    change_id: str
    accepted: bool
    note: str = ""


@api_router.post("/meetings/{meeting_id}/agent/feedback")
async def agent_feedback(
    meeting_id: str,
    payload: AgentFeedback,
    user=Depends(get_current_user),
):
    change = await db.agent_changes.find_one(
        {"id": payload.change_id, "meeting_id": meeting_id, "user_id": user["id"]}
    )
    if not change:
        raise HTTPException(status_code=404, detail="Alteração não encontrada")

    change = agent_svc.record_feedback(change, payload.accepted, payload.note)

    if payload.accepted:
        meeting = await db.meetings.find_one({"id": meeting_id, "user_id": user["id"]}, {"_id": 0})
        if meeting:
            meeting = agent_svc.apply_change_to_meeting(meeting, change)
            await db.meetings.update_one(
                {"id": meeting_id},
                {"$set": {"ata_content": meeting.get("ata_content", ""), "agenda": meeting.get("agenda", []), "updated_at": now_iso()}},
            )

    await db.agent_changes.update_one(
        {"id": payload.change_id},
        {"$set": {"status": change["status"], "feedback_at": change.get("feedback_at"), "user_note": change.get("user_note", "")}},
    )
    return {"ok": True, "status": change["status"]}


@api_router.get("/agent/profile")
async def agent_profile(user=Depends(get_current_user)):
    accepted = await db.agent_changes.count_documents({"user_id": user["id"], "status": "accepted"})
    rejected = await db.agent_changes.count_documents({"user_id": user["id"], "status": "rejected"})
    total = accepted + rejected
    kb_count = await db.knowledge_docs.count_documents({"user_id": user["id"]})
    acceptance_rate = round(accepted / total * 100) if total > 0 else 0
    maturity = "alta" if total > 50 else ("média" if total > 10 else "baixa")
    return {
        "total_feedbacks": total,
        "accepted_changes": accepted,
        "rejected_changes": rejected,
        "acceptance_rate": acceptance_rate,
        "kb_documents": kb_count,
        "maturity": maturity,
        "learning_active": kb_count > 0 or total > 0,
    }


# =========================
# Economic Data
# =========================
@api_router.get("/economicos")
async def economic_data(user=Depends(get_current_user)):
    indicators = await get_economic_indicators()
    selic = indicators.get("selic", {}).get("value", 13.75)
    ipca = indicators.get("ipca", {}).get("value", 4.83)
    scenario = generate_scenario_analysis(selic, ipca)
    return {**indicators, "scenario": scenario}


@api_router.get("/economicos/historico/selic")
async def selic_history(months: int = 12, user=Depends(get_current_user)):
    return await get_selic_history(months)


@api_router.get("/economicos/historico/ipca")
async def ipca_history(months: int = 12, user=Depends(get_current_user)):
    return await get_ipca_history(months)


# =========================
# Admin Settings
# =========================
@api_router.get("/admin/settings")
async def get_admin_settings(user=Depends(get_current_user)):
    doc = await db.admin_settings.find_one({"user_id": user["id"]}, {"_id": 0})
    if not doc:
        return {
            "user_id": user["id"],
            "institution_name": user.get("institution", ""),
            "ata_style": "rpps_padrao",
            "formal_markers": [],
            "custom_opening": "",
            "custom_closing": "",
            "logo_url": "",
            "brasao_url": "",
            "default_president": "",
            "default_secretary": "",
        }
    return doc


class AdminSettings(BaseModel):
    institution_name: Optional[str] = None
    ata_style: Optional[str] = None
    formal_markers: Optional[List[str]] = None
    custom_opening: Optional[str] = None
    custom_closing: Optional[str] = None
    logo_url: Optional[str] = None
    brasao_url: Optional[str] = None
    default_president: Optional[str] = None
    default_secretary: Optional[str] = None


@api_router.put("/admin/settings")
async def update_admin_settings(payload: AdminSettings, user=Depends(get_current_user)):
    update = {k: v for k, v in payload.model_dump(exclude_none=True).items()}
    update["user_id"] = user["id"]
    update["updated_at"] = now_iso()
    await db.admin_settings.update_one(
        {"user_id": user["id"]}, {"$set": update}, upsert=True
    )
    doc = await db.admin_settings.find_one({"user_id": user["id"]}, {"_id": 0})
    return doc


@api_router.get("/admin/users")
async def list_users(user=Depends(get_current_user)):
    if user.get("role") != "admin":
        raise HTTPException(status_code=403, detail="Acesso negado")
    docs = await db.users.find({}, {"_id": 0, "password": 0}).to_list(200)
    return docs


class AdminUserCreate(BaseModel):
    name: str
    email: EmailStr
    password: str
    institution: Optional[str] = ""
    role: Optional[str] = "user"


@api_router.post("/admin/users")
async def create_user(payload: AdminUserCreate, user=Depends(get_current_user)):
    if user.get("role") != "admin":
        raise HTTPException(status_code=403, detail="Acesso negado")
    existing = await db.users.find_one({"email": payload.email.lower()})
    if existing:
        raise HTTPException(status_code=400, detail="E-mail já cadastrado")
    new_user = {
        "id": str(uuid.uuid4()),
        "name": payload.name,
        "email": payload.email.lower(),
        "password": hash_pw(payload.password),
        "institution": payload.institution or "",
        "role": payload.role or "user",
        "created_at": now_iso(),
    }
    await db.users.insert_one(new_user)
    new_user.pop("password")
    new_user.pop("_id", None)
    return new_user


@api_router.delete("/admin/users/{target_id}")
async def delete_user(target_id: str, user=Depends(get_current_user)):
    if user.get("role") != "admin":
        raise HTTPException(status_code=403, detail="Acesso negado")
    if target_id == user["id"]:
        raise HTTPException(status_code=400, detail="Não é possível excluir o próprio usuário")
    await db.users.delete_one({"id": target_id})
    return {"ok": True}


# =========================
# Quality Scoring
# =========================
@api_router.get("/meetings/{meeting_id}/quality")
async def meeting_quality(meeting_id: str, user=Depends(get_current_user)):
    meeting = await db.meetings.find_one({"id": meeting_id, "user_id": user["id"]}, {"_id": 0})
    if not meeting:
        raise HTTPException(status_code=404, detail="Reunião não encontrada")
    scores = score_meeting(meeting)
    gaps = detect_gaps(meeting)
    return {"scores": scores, "gaps": gaps}


# =========================
# Mount router + middleware
# =========================
@api_router.get("/")
async def root():
    return {"app": "Sistema de Atas IA", "version": "2.0"}


app.include_router(api_router)

app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=os.environ.get('CORS_ORIGINS', '*').split(','),
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.on_event("shutdown")
async def shutdown_db_client():
    client.close()


# =========================
# SPA — serve React build
# =========================
_static = ROOT_DIR / "static"
if _static.is_dir():
    @app.get("/{full_path:path}", include_in_schema=False)
    async def _serve_spa(full_path: str):
        f = _static / full_path
        if f.is_file():
            return FileResponse(f)
        return FileResponse(_static / "index.html")
